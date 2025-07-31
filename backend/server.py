import os
import requests
import pandas as pd
from datetime import datetime
from dotenv import load_dotenv
from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load environment variables
dotenv_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), '.env')
load_dotenv(dotenv_path, override=True)

app = Flask(__name__)

# --- Helpers ---

def get_greeting():
    hour = datetime.now().hour
    if hour < 12:
        return "Good Morning"
    elif hour < 17:
        return "Good Afternoon"
    return "Good Evening"

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            element = OxmlElement(f'w:{edge}')
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

# --- API Routes ---

@app.route("/")
def index():
    return jsonify({"message": "Invoice Assistant API is up!"})


@app.route("/read_timesheet", methods=["GET"])
def read_timesheet():
    filename = request.args.get("filename")
    if not filename:
        return jsonify({"error": "Missing filename"}), 400

    try:
        filepath = os.path.join(os.getcwd(), filename)
        if not os.path.exists(filepath):
            df = pd.DataFrame(columns=["Date", "Status", "Remarks"])
            df.to_excel(filepath, index=False)

        df = pd.read_excel(filepath)
        if df.empty:
            return jsonify({"result": "The Excel file is empty."})

        lines = [f"{row['Date']} | {row['Status']} | {row['Remarks']}" for _, row in df.iterrows()]
        return jsonify({"result": "Timesheet Records:\n" + "\n".join(lines)})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/save_or_update_timesheet", methods=["POST"])
def save_or_update_timesheet():
    data = request.json
    filename = data.get("filename")
    date = data.get("date")
    status = data.get("status")
    remarks = data.get("remarks")

    if not all([filename, date, status, remarks]):
        return jsonify({"error": "Missing required fields"}), 400

    try:
        filepath = os.path.join(os.getcwd(), filename)
        columns = ["Date", "Status", "Remarks"]

        if os.path.exists(filepath):
            df = pd.read_excel(filepath)
            df["Date"] = df["Date"].astype(str).str.split().str[0]
        else:
            df = pd.DataFrame(columns=columns)

        if date in df["Date"].values:
            idx = df.index[df["Date"] == date][0]
            df.loc[idx, "Status"] = status
            df.loc[idx, "Remarks"] = remarks
            action = "updated"
        else:
            new_entry = pd.DataFrame([{"Date": date, "Status": status, "Remarks": remarks}])
            df = pd.concat([df, new_entry], ignore_index=True)
            action = "added"

        df.to_excel(filepath, index=False)
        return jsonify({"result": f"Success: Entry for {date} was {action} in {filename}."})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/create_invoice", methods=["POST"])
def create_invoice():
    content = request.json
    filename = content.get("filename")
    data = content.get("data")

    if not filename or not data:
        return jsonify({"error": "Missing filename or data"}), 400

    if not filename.endswith(".docx"):
        filename += ".docx"

    try:
        file_path = os.path.join(os.getcwd(), filename)
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)

        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(6.5)

        # Row 1: INVOICE
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        run = p.add_run("INVOICE")
        run.font.name = "Arial Black"
        run.font.size = Pt(28)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"}, top={"sz": 6, "val": "single"})

        # Row 2: Name and Date
        row_cell = table.add_row().cells[0]
        inner_table = row_cell.add_table(rows=1, cols=2)
        inner_table.columns[0].width = Inches(4.0)
        inner_table.columns[1].width = Inches(2.5)

        inner_table.cell(0, 0).paragraphs[0].add_run(data["name"]).bold = True
        inner_table.cell(0, 1).paragraphs[0].add_run(data["date"])
        inner_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in inner_table._cells:
            set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
        set_cell_border(row_cell, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"}, top={"val": "nil"})

        # Row 3: Bill To
        cell = table.add_row().cells[0]
        p = cell.paragraphs[0]
        p.add_run("Bill To:\n").bold = True
        for line in data["bill_to"]:
            p.add_run(f"    {line}\n")
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # Row 4: Header
        row_cell = table.add_row().cells[0]
        header_table = row_cell.add_table(rows=1, cols=2)
        header_table.columns[0].width = Inches(5.0)
        header_table.columns[1].width = Inches(1.5)

        desc_cell = header_table.cell(0, 0)
        amt_cell = header_table.cell(0, 1)
        desc_cell.text = "DESCRIPTION"
        amt_cell.text = "AMOUNT"

        for cell in [desc_cell, amt_cell]:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            shade = OxmlElement('w:shd')
            shade.set(qn('w:fill'), "ff99cc")
            cell._tc.get_or_add_tcPr().append(shade)

        set_cell_border(row_cell, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # Row 5: Details
        row_cell = table.add_row().cells[0]
        details_table = row_cell.add_table(rows=0, cols=2)
        details_table.columns[0].width = Inches(5.0)
        details_table.columns[1].width = Inches(1.5)

        cells = details_table.add_row().cells
        cells[0].text = data["salary_description"]
        cells[1].text = ""

        for item in data["details"]:
            cells = details_table.add_row().cells
            if isinstance(item, dict):
                cells[0].text = item.get("description", "")
                cells[1].text = item.get("amount", "")
            elif isinstance(item, str) and ":" in item:
                key, val = item.split(":", 1)
                cells[0].text = key.strip()
                cells[1].text = val.strip()
            else:
                cells[0].text = str(item)
                cells[1].text = ""
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        set_cell_border(row_cell, top={"val": "nil"}, bottom={"sz": 6, "val": "single"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # Row 6: Total
        row_cell = table.add_row().cells[0]
        total_table = row_cell.add_table(rows=1, cols=2)
        total_table.columns[0].width = Inches(5.0)
        total_table.columns[1].width = Inches(1.5)

        total_table.cell(0, 0).paragraphs[0].add_run("TOTAL").bold = True
        total_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_table.cell(0, 1).paragraphs[0].add_run(data["total"]).bold = True
        total_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        shade = OxmlElement('w:shd')
        shade.set(qn('w:fill'), "ffcc99")
        total_table.cell(0, 1)._tc.get_or_add_tcPr().append(shade)

        set_cell_border(row_cell, top={"sz": 6, "val": "single"}, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # Row 7: Amount in words
        cell = table.add_row().cells[0]
        p = cell.paragraphs[0]
        p.add_run("Amount in Words: ").bold = True
        p.add_run(data["total_words"])
        set_cell_border(cell, top={"val": "nil"}, bottom={"sz": 6, "val": "single"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        doc.save(file_path)
        return jsonify({"result": f"Invoice saved to {file_path}."})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/send_telegram", methods=["POST"])
def send_telegram():
    data = request.json
    xlsx = data.get("xlsx_filename")
    docx = data.get("docx_filename")

    try:
        token = os.getenv("TELEGRAM_BOT_TOKEN")
        chat_id = os.getenv("TELEGRAM_CHAT_ID")
        if not token or not chat_id:
            raise ValueError("Missing TELEGRAM_BOT_TOKEN or CHAT_ID")

        greeting = get_greeting()
        month = datetime.now().strftime('%B')
        message = (
            f"Hi,\n{greeting}.\n\n"
            f"I've attached the timesheet and invoice for {month}.\n"
            "Please review and approve."
        )

        requests.post(f"https://api.telegram.org/bot{token}/sendMessage", data={
            "chat_id": chat_id,
            "text": message
        })

        count = 0
        for f in [xlsx, docx]:
            if os.path.exists(f):
                with open(f, 'rb') as file:
                    r = requests.post(
                        f"https://api.telegram.org/bot{token}/sendDocument",
                        data={"chat_id": chat_id},
                        files={"document": (f, file)}
                    )
                    if r.status_code == 200:
                        count += 1

        if count == 0:
            return jsonify({"error": "No files sent."}), 400

        return jsonify({"result": f"{count} file(s) sent to Telegram."})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/calculate_salary", methods=["GET"])
def calculate_salary():
    try:
        present_days = request.args.get("present_days", type=int)
        pay_per_day = request.args.get("pay_per_day", type=int)

        if present_days is None or pay_per_day is None:
            return jsonify({"error": "Missing 'present_days' or 'pay_per_day'"}), 400

        total_salary = present_days * pay_per_day
        return jsonify({
            "present_days": present_days,
            "pay_per_day": pay_per_day,
            "salary": total_salary
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True)
